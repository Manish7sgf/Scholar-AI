"""DOCX export service for generating formatted documents"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict
from io import BytesIO
from app.services.template_service import template_service


class DocxService:
    """Service for exporting papers to DOCX format"""
    
    def create_document(
        self,
        title: str,
        authors: List[str],
        abstract: str,
        keywords: List[str],
        sections: List[Dict[str, str]],
        template_format: str,
        ai_disclosure: str = None
    ) -> BytesIO:
        """
        Create a formatted DOCX document
        
        Args:
            title: Paper title
            authors: List of author names
            abstract: Abstract text
            keywords: List of keywords
            sections: List of sections with 'name' and 'content' keys
            template_format: Template format name (ieee, springer, etc.)
            ai_disclosure: Optional AI disclosure statement
            
        Returns:
            BytesIO: Document as bytes
        """
        # Get template configuration
        template = template_service.get_template(template_format)
        
        # Create document
        doc = Document()
        
        # Set page margins
        sections_obj = doc.sections
        for section in sections_obj:
            section.top_margin = Inches(template.margin_top)
            section.bottom_margin = Inches(template.margin_bottom)
            section.left_margin = Inches(template.margin_left)
            section.right_margin = Inches(template.margin_right)
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(template.font_size + 4)
        title_run.font.name = template.font_family
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add authors
        if authors:
            authors_para = doc.add_paragraph()
            authors_run = authors_para.add_run(", ".join(authors))
            authors_run.font.size = Pt(template.font_size + 1)
            authors_run.font.name = template.font_family
            authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Add abstract
        if abstract:
            abstract_heading = doc.add_heading("Abstract", level=1)
            abstract_heading.style.font.size = Pt(template.font_size + 2)
            abstract_heading.style.font.name = template.font_family
            
            abstract_para = doc.add_paragraph(abstract)
            abstract_para.style.font.size = Pt(template.font_size)
            abstract_para.style.font.name = template.font_family
            
            doc.add_paragraph()  # Empty line
        
        # Add keywords
        if keywords:
            keywords_para = doc.add_paragraph()
            keywords_label = keywords_para.add_run("Keywords: ")
            keywords_label.bold = True
            keywords_label.font.size = Pt(template.font_size)
            keywords_label.font.name = template.font_family
            
            keywords_text = keywords_para.add_run(", ".join(keywords))
            keywords_text.font.size = Pt(template.font_size)
            keywords_text.font.name = template.font_family
            
            doc.add_paragraph()  # Empty line
        
        # Add sections
        for section in sections:
            section_name = section.get("name", "")
            section_content = section.get("content", "")
            
            if section_name and section_content:
                # Add section heading
                heading = doc.add_heading(section_name, level=1)
                heading.style.font.size = Pt(template.font_size + 2)
                heading.style.font.name = template.font_family
                
                # Add section content
                content_para = doc.add_paragraph(section_content)
                content_para.style.font.size = Pt(template.font_size)
                content_para.style.font.name = template.font_family
                
                doc.add_paragraph()  # Empty line
        
        # Add AI disclosure if provided
        if ai_disclosure:
            doc.add_page_break()
            disclosure_heading = doc.add_heading("AI Disclosure", level=1)
            disclosure_heading.style.font.size = Pt(template.font_size + 2)
            disclosure_heading.style.font.name = template.font_family
            
            disclosure_para = doc.add_paragraph(ai_disclosure)
            disclosure_para.style.font.size = Pt(template.font_size)
            disclosure_para.style.font.name = template.font_family
        
        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream


# Singleton instance
docx_service = DocxService()
